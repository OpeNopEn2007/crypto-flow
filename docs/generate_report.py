#!/usr/bin/env python3
"""生成 CryptoFlow 大作业实验报告 (docx)"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ── 全局样式 ──
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '黑体'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    if level == 1:
        hs.font.size = Pt(18)
    elif level == 2:
        hs.font.size = Pt(15)
    else:
        hs.font.size = Pt(13)

# ── 封面 ──
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('南开大学 计算机大类')
r.font.size = Pt(16)
r.font.name = '黑体'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('高级语言程序设计')
r.font.size = Pt(22)
r.font.name = '黑体'
r.bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('实验报告')
r.font.size = Pt(22)
r.font.name = '黑体'
r.bold = True

for _ in range(4):
    doc.add_paragraph()

info_items = ['姓名：', '学号：', '班级：', '2026 年 5 月']
for item in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(item)
    r.font.size = Pt(14)
    r.font.name = '宋体'

doc.add_page_break()

# ── 目录页 ──
doc.add_heading('目录', level=1)
toc_items = [
    '一、作业题目',
    '二、开发软件',
    '三、课题要求',
    '四、主要流程',
    '    4.1 整体架构',
    '    4.2 算法原理',
    '        4.2.1 凯撒密码',
    '        4.2.2 维吉尼亚密码',
    '        4.2.3 RSA',
    '        4.2.4 Base64 编码',
    '        4.2.5 XOR 加密',
    '    4.3 动画系统设计',
    '五、单元测试',
    '六、AI 工具使用说明',
    '七、收获',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ═══════════════════════════════════════
# 一、作业题目
# ═══════════════════════════════════════
doc.add_heading('一、作业题目', level=1)
doc.add_paragraph('交互式密码学原理可视化演示系统（CryptoFlow）')
doc.add_paragraph(
    '本项目是一个基于 Qt 6 的图形化密码学教学工具，通过交互式动画演示 5 种经典密码算法的加密/解密过程，'
    '帮助用户直观理解密码学原理。支持凯撒密码、维吉尼亚密码、RSA、Base64 编码和 XOR 加密。'
)

# ═══════════════════════════════════════
# 二、开发软件
# ═══════════════════════════════════════
doc.add_heading('二、开发软件', level=1)

table = doc.add_table(rows=6, cols=2, style='Table Grid')
table.columns[0].width = Cm(5)
table.columns[1].width = Cm(11)
data = [
    ('项目', '说明'),
    ('编程语言', 'C++20'),
    ('GUI 框架', 'Qt 6 (Widgets)'),
    ('构建系统', 'CMake 3.20+'),
    ('开发环境', 'macOS + VS Code / Qt Creator'),
    ('版本管理', 'Git + GitHub'),
]
for i, (k, v) in enumerate(data):
    table.cell(i, 0).text = k
    table.cell(i, 1).text = v
    if i == 0:
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

# ═══════════════════════════════════════
# 三、课题要求
# ═══════════════════════════════════════
doc.add_heading('三、课题要求', level=1)
doc.add_paragraph('本课题要求使用 C++ 语言完成一个图形化小程序，具体要求如下：')

requirements = [
    '使用 C++ 图形化平台（本项目选用 Qt 6）实现图形界面。',
    '程序内容具有创新性和实用性，能够直观展示密码学算法的工作原理。',
    '支持多种经典密码算法的可视化演示，用户可交互式地调整参数并观察加密过程。',
    '代码结构清晰，采用面向对象设计，算法层与界面层解耦。',
    '支持命令行自动演示模式，方便批量化展示。',
]
for req in requirements:
    doc.add_paragraph(req, style='List Bullet')

# ═══════════════════════════════════════
# 四、主要流程
# ═══════════════════════════════════════
doc.add_heading('四、主要流程', level=1)

# 4.1 整体架构
doc.add_heading('4.1 整体架构', level=2)
doc.add_paragraph(
    '项目采用经典的 MVC 分层架构，将算法逻辑、可视化场景和用户控制面板分离：'
)

arch_items = [
    '算法层 (src/crypto/)：纯计算逻辑，包含 Caesar、RSA、Vigenere 三个算法类，与 UI 完全解耦。',
    '场景层 (src/scenes/)：继承 QGraphicsScene，负责动画渲染。每个算法对应一个场景类，'
    '包括 CaesarScene、RSAScene、VigenereScene、Base64Scene、XORScene。',
    '控制面板 (src/controlpanel.h/cpp)：QStackedWidget 动态切换 5 种算法的参数控件，'
    '包含输入校验和速度调节。',
    '主窗口 (src/mainwindow.h/cpp)：QSplitter 左(3):右(1) 布局，左侧为动画场景，右侧为控制面板。',
]
for item in arch_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('整体数据流：')
flow_steps = [
    '用户在控制面板选择算法并输入参数',
    '控制面板校验输入合法性（素数检查、范围检查等）',
    '调用算法层进行加密/解密计算',
    '场景层根据计算步骤启动动画序列',
    '动画完成后自动截图保存',
]
for i, step in enumerate(flow_steps, 1):
    doc.add_paragraph(f'{i}. {step}')

# 4.2 算法原理
doc.add_heading('4.2 算法原理', level=2)

# 4.2.1 凯撒密码
doc.add_heading('4.2.1 凯撒密码', level=3)
doc.add_paragraph(
    '凯撒密码是最经典的替换密码，通过将字母表中每个字母替换为固定偏移量后的字母实现加密。'
)
doc.add_paragraph('加密公式：')
doc.add_paragraph('    C = (P + shift) mod 26')
doc.add_paragraph('解密公式：')
doc.add_paragraph('    P = (C - shift) mod 26')
doc.add_paragraph('其中 P 为明文字母对应的数字（A=0, B=1, ...），C 为密文数字，shift 为偏移量。')
doc.add_paragraph()
doc.add_paragraph(
    '可视化实现：采用双圈转盘设计。外圈固定显示标准字母表（明文），内圈显示偏移后的字母表（密文）。'
    '动画时内圈逆时针旋转对齐到指定偏移量，然后逐字母高亮显示加密映射关系，'
    '同时用连接线和脉冲发光效果展示对应关系。'
)

# 4.2.2 维吉尼亚密码
doc.add_heading('4.2.2 维吉尼亚密码', level=3)
doc.add_paragraph(
    '维吉尼亚密码是凯撒密码的多表扩展，使用一个关键词循环作为偏移量序列，'
    '不同位置的字母使用不同的偏移量，大幅提高了密码强度。'
)
doc.add_paragraph('加密公式：')
doc.add_paragraph('    C[i] = (P[i] + K[i mod len(K)]) mod 26')
doc.add_paragraph('其中 K 为关键词字母序列，len(K) 为关键词长度。')
doc.add_paragraph()
doc.add_paragraph(
    '可视化实现：同样使用双圈转盘，但内圈偏移量会根据关键词的每个字母动态切换。'
    '关键词以黄色高亮显示在顶部，每处理一个明文字母，内圈旋转到对应密钥字母的偏移位置，'
    '处理完成后回转复位再处理下一个。密钥索引只对字母字符递增，非字母（空格、标点）不消耗密钥位。'
)

# 4.2.3 RSA
doc.add_heading('4.2.3 RSA', level=3)
doc.add_paragraph(
    'RSA 是最著名的非对称加密算法，基于大整数分解的数学难题。其安全性依赖于：'
    '将两个大素数相乘很容易，但将乘积分解回两个素数极其困难。'
)
doc.add_paragraph('密钥生成步骤：')
rsa_steps = [
    '选择两个不相等的大素数 p 和 q',
    '计算 n = p × q（模数）',
    '计算欧拉函数 φ(n) = (p-1)(q-1)',
    '选择公钥指数 e，满足 1 < e < φ(n) 且 gcd(e, φ(n)) = 1',
    '计算私钥指数 d，满足 d × e ≡ 1 (mod φ(n))，即 d = e⁻¹ mod φ(n)',
    '公钥 = (e, n)，私钥 = (d, n)',
]
for i, step in enumerate(rsa_steps, 1):
    doc.add_paragraph(f'({i}) {step}')

doc.add_paragraph()
doc.add_paragraph('加密：C = M^e mod n')
doc.add_paragraph('解密：M = C^d mod n')
doc.add_paragraph()
doc.add_paragraph(
    '可视化实现：采用 7 步卡片流程动画，每一步展示密钥生成的一个阶段。'
    '居中卡片布局，圆形步骤编号，关键数值以绿色高亮显示。'
    '加密后自动将密文存入缓存，切换到解密模式时自动填充，方便用户理解加解密的对称性。'
    '模幂运算使用 __int128 中间变量防止 64 位整数乘法溢出。'
)

# 4.2.4 Base64
doc.add_heading('4.2.4 Base64 编码', level=3)
doc.add_paragraph(
    'Base64 是一种将二进制数据编码为 ASCII 字符的编码方式，'
    '常用于在仅支持文本的传输通道中安全地传递二进制数据。'
)
doc.add_paragraph('编码过程：')
base64_steps = [
    '将输入文本转换为 ASCII 码的二进制表示',
    '每 8 位一组重新划分为每 6 位一组',
    '将每组 6 位二进制转换为十进制数（0-63）',
    '按照 Base64 编码表映射为对应的 ASCII 字符',
    '输出长度不是 4 的倍数时，用 "=" 填充',
]
for i, step in enumerate(base64_steps, 1):
    doc.add_paragraph(f'({i}) {step}')

doc.add_paragraph()
doc.add_paragraph(
    '可视化实现：5 步卡片流程，展示明文 → 二进制 → 6bit 分组 → 查表 → 最终结果的完整转换过程，'
    '与 RSA 场景共享卡片布局风格。'
)

# 4.2.5 XOR
doc.add_heading('4.2.5 XOR 加密', level=3)
doc.add_paragraph(
    'XOR（异或）加密是最简单的对称加密算法，利用异或运算的自反性：'
    'A ⊕ K = C，C ⊕ K = A。加密和解密使用完全相同的操作。'
)
doc.add_paragraph('逐字节运算：')
doc.add_paragraph('    密文字节 = 明文字节 ⊕ 密钥字节')
doc.add_paragraph()
doc.add_paragraph(
    '可视化实现：4 步卡片流程，分别展示明文、密钥、逐字节 XOR 运算过程和加密结果。'
    '同时以二进制形式展示每一字节的异或运算细节。'
)

# 4.3 动画系统设计
doc.add_heading('4.3 动画系统设计', level=2)
doc.add_paragraph(
    '动画系统是本项目的核心难点之一。由于 QGraphicsItem 不支持 QPropertyAnimation，'
    '因此采用 QTimer + QEasingCurve 的手动动画方案：'
)

anim_items = [
    '离散步进旋转：每步 500ms 暂停 + 12 帧 OutCubic 缓动旋转，让用户看清每一步变化。',
    'animationId_ 防并发：每次启动动画时递增 ID，所有 QTimer::singleShot 回调执行前校验 ID，'
    '过期则丢弃，避免旧动画回调干扰新动画。',
    '回转复位机制：旋转完成后匀速回转初始位置（30 帧线性），停顿 500ms 后再开始逐字母高亮。',
    '内圈字母不使用 QGraphicsItemGroup（旋转会导致字母倾斜），而是手动用 cos/sin 重新计算每个字母的位置。',
    '脉冲发光效果：用 QGraphicsEllipseItem + QLinearGradient 实现连接线端点的发光动画。',
]
for item in anim_items:
    doc.add_paragraph(item, style='List Bullet')

# ═══════════════════════════════════════
# 五、单元测试
# ═══════════════════════════════════════
doc.add_heading('五、单元测试', level=1)
doc.add_paragraph(
    '本项目的算法层（crypto/）设计为纯计算逻辑，与 UI 完全解耦，便于独立测试。'
    '以下是各算法的关键测试用例：'
)

doc.add_paragraph()
doc.add_heading('凯撒密码测试', level=3)
t = doc.add_table(rows=4, cols=3, style='Table Grid')
for i, row_data in enumerate([
    ('输入', '偏移量', '预期输出'),
    ('HELLO', '3', 'KHOOR'),
    ('WORLD', '5', 'BTWQI'),
    ('ABC', '25', 'ZAB'),
]):
    for j, val in enumerate(row_data):
        t.cell(i, j).text = val

doc.add_paragraph()
doc.add_heading('RSA 测试', level=3)
t = doc.add_table(rows=4, cols=3, style='Table Grid')
for i, row_data in enumerate([
    ('p', 'q', '验证'),
    ('61', '53', 'n=3233, φ=3120, e=17, d=2753'),
    ('7', '11', 'n=77, φ=60, e=13, d=37'),
    ('13', '17', 'n=221, φ=192, e=5, d=77'),
]):
    for j, val in enumerate(row_data):
        t.cell(i, j).text = val

doc.add_paragraph()
doc.add_heading('维吉尼亚密码测试', level=3)
t = doc.add_table(rows=3, cols=3, style='Table Grid')
for i, row_data in enumerate([
    ('明文', '密钥', '密文'),
    ('HELLO', 'KEY', 'RIJVS'),
    ('ATTACK', 'LEMON', 'LXFOPV'),
]):
    for j, val in enumerate(row_data):
        t.cell(i, j).text = val

doc.add_paragraph()
doc.add_paragraph(
    '此外，ControlPanel 中对所有算法的输入都进行了校验：'
)
validations = [
    'RSA：p、q 必须为素数且不相等，p、q >= 3',
    'RSA 加密：明文 M 必须满足 0 ≤ M < n',
    '维吉尼亚/XOR：密钥和输入不能为空',
    'Base64：输入不能为空',
    'CLI --shift 越界时默认回退到 3',
]
for v in validations:
    doc.add_paragraph(v, style='List Bullet')

# ═══════════════════════════════════════
# 六、AI 工具使用说明
# ═══════════════════════════════════════
doc.add_heading('六、AI 工具使用说明', level=1)
doc.add_paragraph('本项目开发过程中使用了以下 AI 工具：')

ai_table = doc.add_table(rows=4, cols=3, style='Table Grid')
for i, row_data in enumerate([
    ('工具名称', '用途', '使用位置'),
    ('Claude (Anthropic)', '代码开发辅助、架构设计、Bug 调试', '全部源代码'),
    ('Claude Code (CLI)', '自动化代码生成、重构、测试', 'src/ 目录下所有文件'),
    ('GitHub Copilot', '代码补全', '部分代码编写'),
]):
    for j, val in enumerate(row_data):
        ai_table.cell(i, j).text = val
        if i == 0:
            for paragraph in ai_table.cell(i, j).paragraphs:
                for run in paragraph.runs:
                    run.bold = True

doc.add_paragraph()
doc.add_paragraph(
    '说明：项目从零开始由 AI 辅助完成。算法设计思路、Qt 动画方案、架构分层均由开发者主导，'
    'AI 主要用于代码实现、调试和优化。所有代码经过人工审查和实际运行验证。'
)

# ═══════════════════════════════════════
# 七、收获
# ═══════════════════════════════════════
doc.add_heading('七、收获', level=1)

doc.add_heading('7.1 Qt 6 图形编程', level=2)
doc.add_paragraph(
    '通过本项目深入学习了 Qt 6 的 QGraphicsView/QGraphicsScene 框架。'
    '理解了 Qt 的坐标系统、信号槽机制、QTimer 定时器的使用方式。'
    '特别是 QGraphicsItem 不支持属性动画这一限制，'
    '学会了用 QTimer + QEasingCurve 手动实现平滑动画的方案。'
)

doc.add_heading('7.2 面向对象设计', level=2)
doc.add_paragraph(
    '项目严格遵循面向对象原则：算法层与视图层分离，每个密码算法都有独立的算法类和场景类。'
    '这种分层设计使得添加新算法只需实现对应的 crypto 和 scenes 类，'
    '无需修改主窗口和控制面板的核心逻辑。'
)

doc.add_heading('7.3 动画系统设计', level=2)
doc.add_paragraph(
    '动画系统的并发控制是一个重要收获。animationId_ 机制解决了多个动画回调互相干扰的问题，'
    '回转复位机制确保场景状态的正确性。这些设计模式在后续的图形编程中都会用到。'
)

doc.add_heading('7.4 密码学原理', level=2)
doc.add_paragraph(
    '通过可视化实现加深了对密码学原理的理解。'
    '凯撒密码和维吉尼亚密码让我理解了古典密码的替换思想及其局限性；'
    'RSA 算法让我接触了模运算、欧拉函数、模逆元等数论知识；'
    'Base64 编码让我理解了二进制数据与文本之间的转换方式。'
)

doc.add_heading('7.5 C++ 工程实践', level=2)
doc.add_paragraph(
    '学习了 CMake 构建系统的配置、Qt 的 MOC/UIC/RCC 机制、'
    '跨平台编译的注意事项（如 MSVC 对 __int128 的兼容处理）。'
    '也学会了使用 Git 进行版本管理、GitHub Actions 进行 CI/CD 自动化构建。'
)

# ── 保存 ──
out_dir = os.path.dirname(os.path.abspath(__file__))
out_path = os.path.join(out_dir, 'CryptoFlow_实验报告.docx')
doc.save(out_path)
print(f'报告已生成: {out_path}')
